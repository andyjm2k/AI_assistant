"""Tests for CATBot modular skills framework."""

from __future__ import annotations

import base64
import json
import shutil
import uuid
from pathlib import Path
from typing import Any, Dict, Sequence
from unittest.mock import AsyncMock, MagicMock, patch

import openpyxl
import pytest
from docx import Document
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from src.skills import BaseSkill, BaseTool, SkillContext, SkillManager, SkillValidationError


class EchoTool(BaseTool):
    name = "echo"
    description = "Echo test tool."
    input_schema = {"type": "object", "properties": {"text": {"type": "string"}}}

    async def run(self, arguments: Dict[str, Any], context: SkillContext) -> Dict[str, Any]:
        return {"text": str(arguments.get("text", ""))}


class AnotherEchoTool(BaseTool):
    name = "echo"
    description = "Second echo tool to force alias ambiguity."
    input_schema = {"type": "object", "properties": {"text": {"type": "string"}}}

    async def run(self, arguments: Dict[str, Any], context: SkillContext) -> Dict[str, Any]:
        return {"text": f"secondary:{arguments.get('text', '')}"}


class AlphaSkill(BaseSkill):
    name = "alpha"
    description = "alpha"

    def create_tools(self) -> Sequence[BaseTool]:
        return [EchoTool()]


class BetaSkill(BaseSkill):
    name = "beta"
    description = "beta"

    def create_tools(self) -> Sequence[BaseTool]:
        return [AnotherEchoTool()]


def test_manifest_loading_discovers_builtin_skills() -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    names = [spec.name for spec in manager.list_skills()]
    assert "core" in names
    assert "filesystem" in names
    assert "GitHubProjectManager" in names
    assert "googleworkspace_cli" in names
    assert "image_generation" in names
    assert "spotify_player" in names
    assert "telegram_admin" in names
    assert "testkit" in names

    tool_names = [spec.qualified_name for spec in manager.list_tools()]
    assert "core.ping" in tool_names
    assert "core.echo" in tool_names
    assert "filesystem.list_files" in tool_names
    assert "filesystem.search_files" in tool_names
    assert "GitHubProjectManager.status" in tool_names
    assert "GitHubProjectManager.fetch" in tool_names
    assert "GitHubProjectManager.sync" in tool_names
    assert "GitHubProjectManager.list_pull_requests" in tool_names
    assert "googleworkspace_cli.check_cli" in tool_names
    assert "googleworkspace_cli.check_auth" in tool_names
    assert "googleworkspace_cli.gmail_list_unread" in tool_names
    assert "googleworkspace_cli.gmail_list_all" in tool_names
    assert "googleworkspace_cli.gmail_get_message" in tool_names
    assert "googleworkspace_cli.gmail_compose_draft" in tool_names
    assert "googleworkspace_cli.gmail_send_message" in tool_names
    assert "googleworkspace_cli.gmail_mark_read" in tool_names
    assert "googleworkspace_cli.slides_batch_update_presentation" in tool_names
    assert "googleworkspace_cli.slides_create_presentation_from_markdown" in tool_names
    assert "googleworkspace_cli.list_available_commands" in tool_names
    assert "googleworkspace_cli.run_readonly_command" in tool_names
    assert "image_generation.generate_image" in tool_names
    assert "spotify_player.search_tracks" in tool_names
    assert "spotify_player.get_available_devices" in tool_names
    assert "spotify_player.play_track" in tool_names
    assert "spotify_player.play_playlist" in tool_names
    assert "telegram_admin.notify_admin" in tool_names
    assert "testkit.context_snapshot" in tool_names

    mcp_tools = manager.mcp_tools()
    assert any(t.get("name") == "core.ping" for t in mcp_tools)


@pytest.mark.asyncio
async def test_execute_core_ping_tool() -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool("core.ping", {})
    assert result.success is True
    assert isinstance(result.data, dict)
    assert result.data.get("pong") is True


@pytest.mark.asyncio
async def test_filesystem_write_and_read_roundtrip() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        manager = SkillManager()
        manager.load_manifests("src/skills/manifests")
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=temp_base / "root"),
            replace=True,
        )

        write_result = await manager.execute_tool(
            "filesystem.write_text",
            {"path": "notes/test.txt", "content": "hello"},
        )
        assert write_result.success is True

        read_result = await manager.execute_tool("filesystem.read_text", {"path": "notes/test.txt"})
        assert read_result.success is True
        assert read_result.data["content"] == "hello"
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_read_text_supports_line_ranges() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        (root / "notes.txt").write_text("one\ntwo\nthree\nfour\n", encoding="utf-8")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool(
            "filesystem.read_text",
            {"path": "notes.txt", "start_line": 2, "end_line": 3, "include_line_numbers": True},
        )
        assert result.success is True
        assert result.data["content"] == "2: two\n3: three"
        assert result.data["excerpt_start_line"] == 2
        assert result.data["excerpt_end_line"] == 3
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_read_text_supports_docx() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("Quarterly summary")
        doc.add_paragraph("Revenue increased 12 percent")
        doc.save(root / "summary.docx")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool("filesystem.read_text", {"path": "summary.docx"})
        assert result.success is True
        assert result.data["type"] == "text"
        assert "Quarterly summary" in result.data["content"]
        assert "Revenue increased 12 percent" in result.data["content"]
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_read_text_supports_pdf() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        pdf_path = root / "report.pdf"
        pdf_canvas = canvas.Canvas(str(pdf_path), pagesize=letter)
        pdf_canvas.drawString(100, 750, "Quarterly pipeline report")
        pdf_canvas.save()

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool("filesystem.read_text", {"path": "report.pdf"})
        assert result.success is True
        assert result.data["type"] == "text"
        assert "Quarterly pipeline report" in result.data["content"]
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_read_text_supports_images() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        image_path = root / "diagram.png"
        Image.new("RGB", (64, 32), color="red").save(image_path)

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool("filesystem.read_text", {"path": "diagram.png"})
        assert result.success is True
        assert result.data["type"] == "image"
        assert "64x32" in result.data["content"]
        assert result.data["image_data"]["metadata"]["width"] == 64
        assert result.data["image_data"]["metadata"]["height"] == 32
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_list_files_supports_recursive_subdirectories() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        (root / "nested" / "deeper").mkdir(parents=True, exist_ok=True)
        (root / "top.txt").write_text("top", encoding="utf-8")
        (root / "nested" / "child.txt").write_text("child", encoding="utf-8")
        (root / "nested" / "deeper" / "leaf.txt").write_text("leaf", encoding="utf-8")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        non_recursive = await manager.execute_tool("filesystem.list_files", {})
        assert non_recursive.success is True
        non_recursive_paths = {
            str(item.get("relative_path"))
            for item in (non_recursive.data or {}).get("items", [])
            if isinstance(item, dict)
        }
        assert "top.txt" in non_recursive_paths
        assert "nested/child.txt" not in non_recursive_paths

        recursive = await manager.execute_tool("filesystem.list_files", {"recursive": True})
        assert recursive.success is True
        recursive_paths = {
            str(item.get("relative_path"))
            for item in (recursive.data or {}).get("items", [])
            if isinstance(item, dict)
        }
        assert "top.txt" in recursive_paths
        assert "nested/child.txt" in recursive_paths
        assert "nested/deeper/leaf.txt" in recursive_paths
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_search_files_finds_filename_and_content_matches() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        (root / "docs").mkdir(parents=True, exist_ok=True)
        (root / "docs" / "roadmap_notes.txt").write_text("alpha topic\nrelease plan\n", encoding="utf-8")
        (root / "docs" / "misc.txt").write_text("contains alpha in body\n", encoding="utf-8")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool(
            "filesystem.search_files",
            {"query": "alpha", "path": "docs"},
        )
        assert result.success is True
        assert result.data["total_matches"] == 2
        paths = [item.get("relative_path") for item in result.data["items"]]
        assert "docs/misc.txt" in paths
        assert "docs/roadmap_notes.txt" in paths
        content_match = next(item for item in result.data["items"] if item["relative_path"] == "docs/misc.txt")
        assert content_match["line_number"] == 1
        assert "alpha" in content_match["excerpt"].lower()
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_search_files_reads_docx_and_xlsx_content() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("Board update")
        doc.add_paragraph("Alpha milestone approved")
        doc.save(root / "board-update.docx")

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet["A1"] = "Milestone"
        sheet["A2"] = "Alpha milestone approved"
        workbook.save(root / "tracker.xlsx")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool("filesystem.search_files", {"query": "Alpha milestone approved"})
        assert result.success is True
        paths = [item.get("relative_path") for item in result.data["items"]]
        assert "board-update.docx" in paths
        assert "tracker.xlsx" in paths
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_list_files_accepts_scratch_prefixed_path() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "scratch"
        (root / "images").mkdir(parents=True, exist_ok=True)
        (root / "images" / "one.png").write_text("png", encoding="utf-8")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        unix_style = await manager.execute_tool("filesystem.list_files", {"path": "scratch/images"})
        windows_style = await manager.execute_tool("filesystem.list_files", {"path": r"scratch\images"})

        assert unix_style.success is True
        assert windows_style.success is True

        unix_paths = {
            str(item.get("relative_path"))
            for item in (unix_style.data or {}).get("items", [])
            if isinstance(item, dict)
        }
        windows_paths = {
            str(item.get("relative_path"))
            for item in (windows_style.data or {}).get("items", [])
            if isinstance(item, dict)
        }
        assert "images/one.png" in unix_paths
        assert "images/one.png" in windows_paths
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_list_files_supports_offset_pagination() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        root = temp_base / "root"
        root.mkdir(parents=True, exist_ok=True)
        (root / "c.txt").write_text("c", encoding="utf-8")
        (root / "a.txt").write_text("a", encoding="utf-8")
        (root / "b.txt").write_text("b", encoding="utf-8")

        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=root),
            replace=True,
        )

        result = await manager.execute_tool(
            "filesystem.list_files",
            {"offset": 1, "max_entries": 1},
        )
        assert result.success is True
        assert result.data["total_count"] == 3
        assert result.data["returned_count"] == 1
        assert result.data["has_more"] is True
        assert result.data["next_offset"] == 2
        assert [item.get("relative_path") for item in result.data["items"]] == ["b.txt"]
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_filesystem_blocks_path_traversal() -> None:
    temp_base = _create_workspace_temp_dir()
    try:
        manager = SkillManager()
        manager.load_manifests(
            _create_temp_filesystem_manifest_dir(temp_base, root=temp_base / "root"),
            replace=True,
        )
        result = await manager.execute_tool(
            "filesystem.read_text",
            {"path": "../outside.txt"},
        )
        assert result.success is False
        assert result.error_code == "framework_error"
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_ambiguous_unqualified_tool_name_returns_framework_error() -> None:
    manager = SkillManager()
    manager.register_skill(AlphaSkill())
    manager.register_skill(BetaSkill())
    result = await manager.execute_tool("echo", {"text": "hello"})
    assert result.success is False
    assert result.error_code == "framework_error"


@pytest.mark.asyncio
async def test_image_generation_writes_output_file(monkeypatch: pytest.MonkeyPatch) -> None:
    temp_base = _create_workspace_temp_dir()
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    image_bytes = b"fake-png-bytes"
    image_data_url = "data:image/png;base64," + base64.b64encode(image_bytes).decode("utf-8")

    mock_response = MagicMock()
    mock_response.status_code = 200
    mock_response.json.return_value = {
        "id": "img-test-id",
        "provider": "openrouter",
        "choices": [
            {
                "message": {
                    "content": "Generated image",
                    "images": [{"image_url": {"url": image_data_url}}],
                }
            }
        ],
        "usage": {"total_tokens": 5},
    }

    monkeypatch.setenv("IMAGE_GENERATION_OPENROUTER_API_KEY", "test-openrouter-key")
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    try:
        with patch("src.skills.builtin.image_generation_skill.httpx.AsyncClient") as mock_client_cls:
            mock_client = AsyncMock()
            mock_client.__aenter__.return_value = mock_client
            mock_client.__aexit__.return_value = None
            mock_client.post = AsyncMock(return_value=mock_response)
            mock_client_cls.return_value = mock_client

            result = await manager.execute_tool(
                "image_generation.generate_image",
                {"prompt": "A cat astronaut in watercolor", "output_dir": "generated"},
                context=SkillContext(scratch_dir=temp_base),
            )
            post_call = mock_client.post.call_args
            assert post_call is not None
            called_url = post_call.args[0]
            payload = post_call.kwargs["json"]
            assert called_url.endswith("/chat/completions")
            assert payload["model"] == "bytedance-seed/seedream-4.5"
            assert payload["modalities"] == ["image"]
            assert payload["messages"][0]["content"] == "A cat astronaut in watercolor"

        assert result.success is True
        assert result.data["model"] == "bytedance-seed/seedream-4.5"
        assert result.data["image_count"] == 1
        image_item = result.data["images"][0]
        assert image_item["mime_type"] == "image/png"
        assert image_item["relative_path"].startswith("generated/")
        assert Path(image_item["path"]).exists()
        assert Path(image_item["path"]).read_bytes() == image_bytes
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_image_generation_requires_openrouter_api_key(monkeypatch: pytest.MonkeyPatch) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    monkeypatch.delenv("IMAGE_GENERATION_OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("IMAGE_GENERATION_API_KEY", raising=False)
    monkeypatch.delenv("OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("MCP_LLM_OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("openai_api_key", raising=False)
    monkeypatch.delenv("MCP_LLM_OPENAI_API_KEY", raising=False)

    result = await manager.execute_tool(
        "image_generation.generate_image",
        {"prompt": "test prompt"},
    )

    assert result.success is False
    assert result.error_code == "framework_error"


@pytest.mark.asyncio
async def test_image_generation_does_not_fall_back_to_core_openai_key(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    monkeypatch.delenv("IMAGE_GENERATION_OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("IMAGE_GENERATION_API_KEY", raising=False)
    monkeypatch.delenv("OPENROUTER_API_KEY", raising=False)
    monkeypatch.delenv("MCP_LLM_OPENROUTER_API_KEY", raising=False)
    monkeypatch.setenv("OPENAI_API_KEY", "core-openai-key")

    result = await manager.execute_tool(
        "image_generation.generate_image",
        {"prompt": "test prompt"},
    )

    assert result.success is False
    assert result.error_code == "framework_error"


@pytest.mark.asyncio
async def test_spotify_player_search_tracks_uses_client_credentials(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    token_response = MagicMock()
    token_response.status_code = 200
    token_response.json.return_value = {
        "access_token": "spotify-app-token",
        "token_type": "Bearer",
        "expires_in": 3600,
    }

    search_response = MagicMock()
    search_response.status_code = 200
    search_response.json.return_value = {
        "tracks": {
            "total": 1,
            "items": [
                {
                    "id": "0123456789ABCDEFGHIJKL",
                    "name": "CATBot Anthem",
                    "uri": "spotify:track:0123456789ABCDEFGHIJKL",
                    "artists": [
                        {
                            "id": "artist1234567890123456",
                            "name": "CATBot",
                            "uri": "spotify:artist:artist1234567890123456",
                        }
                    ],
                    "album": {
                        "id": "album12345678901234567",
                        "name": "Automation Songs",
                        "release_date": "2026-03-20",
                        "uri": "spotify:album:album12345678901234567",
                    },
                    "duration_ms": 123456,
                    "explicit": False,
                    "popularity": 77,
                    "preview_url": None,
                    "external_urls": {"spotify": "https://open.spotify.com/track/0123456789ABCDEFGHIJKL"},
                    "is_playable": True,
                }
            ],
        }
    }

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.delenv("SPOTIFY_ACCESS_TOKEN", raising=False)
    monkeypatch.delenv("SPOTIFY_REFRESH_TOKEN", raising=False)

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.post = AsyncMock(return_value=token_response)
        mock_client.get = AsyncMock(return_value=search_response)
        mock_client_cls.return_value = mock_client

        result = await manager.execute_tool(
            "spotify_player.search_tracks",
            {"query": "catbot anthem", "limit": 5, "market": "au"},
        )

    assert result.success is True
    assert result.data["query"] == "catbot anthem"
    assert result.data["returned_count"] == 1
    assert result.data["tracks"][0]["name"] == "CATBot Anthem"
    assert result.data["tracks"][0]["artist_names"] == ["CATBot"]

    token_call = mock_client.post.call_args
    assert token_call is not None
    assert token_call.args[0].endswith("/token")
    assert token_call.kwargs["data"] == {"grant_type": "client_credentials"}

    search_call = mock_client.get.call_args
    assert search_call is not None
    assert search_call.args[0].endswith("/search")
    assert search_call.kwargs["params"] == {
        "q": "catbot anthem",
        "type": "track",
        "limit": 5,
        "market": "AU",
    }
    assert search_call.kwargs["headers"]["Authorization"] == "Bearer spotify-app-token"


@pytest.mark.asyncio
async def test_spotify_player_play_track_requires_user_playback_auth(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.delenv("SPOTIFY_ACCESS_TOKEN", raising=False)
    monkeypatch.delenv("SPOTIFY_REFRESH_TOKEN", raising=False)
    monkeypatch.delenv("SPOTIFY_DEVICE_ID", raising=False)

    with pytest.raises(
        SkillValidationError,
        match="SPOTIFY_ACCESS_TOKEN or SPOTIFY_REFRESH_TOKEN",
        ):
            await manager.execute_tool(
                "spotify_player.play_track",
                {"track_id": "0123456789ABCDEFGHIJKL"},
                raise_errors=True,
            )


@pytest.mark.asyncio
async def test_spotify_player_get_available_devices_refreshes_token_and_returns_ids(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    validation_response = MagicMock()
    validation_response.status_code = 401

    token_response = MagicMock()
    token_response.status_code = 200
    token_response.json.return_value = {
        "access_token": "spotify-device-token",
        "token_type": "Bearer",
        "expires_in": 3600,
    }

    devices_response = MagicMock()
    devices_response.status_code = 200
    devices_response.json.return_value = {
        "devices": [
            {
                "id": "device-123",
                "is_active": True,
                "is_private_session": False,
                "is_restricted": False,
                "name": "Office Speaker",
                "type": "Speaker",
                "volume_percent": 42,
                "supports_volume": True,
            },
            {
                "id": "device-456",
                "is_active": False,
                "is_private_session": False,
                "is_restricted": False,
                "name": "Phone",
                "type": "Smartphone",
                "volume_percent": 88,
                "supports_volume": True,
            },
        ]
    }

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.setenv("SPOTIFY_ACCESS_TOKEN", "spotify-expired-token")
    monkeypatch.setenv("SPOTIFY_REFRESH_TOKEN", "spotify-refresh-token")

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.get = AsyncMock(side_effect=[validation_response, devices_response])
        mock_client.post = AsyncMock(return_value=token_response)
        mock_client_cls.return_value = mock_client

        result = await manager.execute_tool(
            "spotify_player.get_available_devices",
            {},
        )

    assert result.success is True
    assert result.data["returned_count"] == 2
    assert result.data["device_ids"] == ["device-123", "device-456"]
    assert result.data["devices"][0]["id"] == "device-123"
    assert result.data["devices"][0]["name"] == "Office Speaker"

    first_get_call = mock_client.get.call_args_list[0]
    assert first_get_call.args[0].endswith("/search")
    assert first_get_call.kwargs["headers"]["Authorization"] == "Bearer spotify-expired-token"

    token_call = mock_client.post.call_args
    assert token_call is not None
    assert token_call.kwargs["data"] == {
        "grant_type": "refresh_token",
        "refresh_token": "spotify-refresh-token",
    }

    devices_call = mock_client.get.call_args_list[1]
    assert devices_call.args[0].endswith("/me/player/devices")
    assert devices_call.kwargs["headers"]["Authorization"] == "Bearer spotify-device-token"


@pytest.mark.asyncio
async def test_spotify_player_play_track_uses_valid_existing_access_token(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    validation_response = MagicMock()
    validation_response.status_code = 200

    play_response = MagicMock()
    play_response.status_code = 204

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.setenv("SPOTIFY_ACCESS_TOKEN", "spotify-existing-token")
    monkeypatch.delenv("SPOTIFY_REFRESH_TOKEN", raising=False)

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.get = AsyncMock(return_value=validation_response)
        mock_client.put = AsyncMock(return_value=play_response)
        mock_client.post = AsyncMock()
        mock_client_cls.return_value = mock_client

        result = await manager.execute_tool(
            "spotify_player.play_track",
            {"track_id": "0123456789ABCDEFGHIJKL"},
        )

    assert result.success is True
    assert result.data["started"] is True
    assert mock_client.post.await_count == 0

    validation_call = mock_client.get.call_args
    assert validation_call is not None
    assert validation_call.args[0].endswith("/search")
    assert validation_call.kwargs["params"] == {
        "q": "spotify",
        "type": "track",
        "limit": 1,
    }
    assert validation_call.kwargs["headers"]["Authorization"] == "Bearer spotify-existing-token"

    play_call = mock_client.put.call_args
    assert play_call is not None
    assert play_call.kwargs["headers"]["Authorization"] == "Bearer spotify-existing-token"


@pytest.mark.asyncio
async def test_spotify_player_play_track_refreshes_expired_access_token(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    validation_response = MagicMock()
    validation_response.status_code = 401

    token_response = MagicMock()
    token_response.status_code = 200
    token_response.json.return_value = {
        "access_token": "spotify-refreshed-token",
        "token_type": "Bearer",
        "expires_in": 3600,
    }

    play_response = MagicMock()
    play_response.status_code = 204

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.setenv("SPOTIFY_ACCESS_TOKEN", "spotify-expired-token")
    monkeypatch.setenv("SPOTIFY_REFRESH_TOKEN", "spotify-refresh-token")

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.get = AsyncMock(return_value=validation_response)
        mock_client.post = AsyncMock(return_value=token_response)
        mock_client.put = AsyncMock(return_value=play_response)
        mock_client_cls.return_value = mock_client

        result = await manager.execute_tool(
            "spotify_player.play_track",
            {"track_id": "0123456789ABCDEFGHIJKL"},
        )

    assert result.success is True
    assert result.data["started"] is True

    validation_call = mock_client.get.call_args
    assert validation_call is not None
    assert validation_call.kwargs["headers"]["Authorization"] == "Bearer spotify-expired-token"

    token_call = mock_client.post.call_args
    assert token_call is not None
    assert token_call.kwargs["data"] == {
        "grant_type": "refresh_token",
        "refresh_token": "spotify-refresh-token",
    }

    play_call = mock_client.put.call_args
    assert play_call is not None
    assert play_call.kwargs["headers"]["Authorization"] == "Bearer spotify-refreshed-token"


@pytest.mark.asyncio
async def test_spotify_player_play_track_invalid_refresh_token_returns_reauth_url(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    token_response = MagicMock()
    token_response.status_code = 400
    token_response.json.return_value = {
        "error": "invalid_grant",
        "error_description": "Invalid refresh token",
    }

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.setenv("SPOTIFY_REFRESH_TOKEN", "spotify-refresh-token")
    monkeypatch.setenv("SPOTIFY_REDIRECT_URI", "https://catbot.local:8002/spotify/callback")
    monkeypatch.delenv("SPOTIFY_ACCESS_TOKEN", raising=False)

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.post = AsyncMock(return_value=token_response)
        mock_client_cls.return_value = mock_client

        with pytest.raises(
            SkillValidationError,
            match=r"one-time URL: https://catbot\.local:8002/spotify/authorize",
        ):
            await manager.execute_tool(
                "spotify_player.play_track",
                {"track_id": "0123456789ABCDEFGHIJKL"},
                raise_errors=True,
            )


@pytest.mark.asyncio
async def test_spotify_player_play_playlist_refreshes_token_and_starts_playback(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    token_response = MagicMock()
    token_response.status_code = 200
    token_response.json.return_value = {
        "access_token": "spotify-user-token",
        "token_type": "Bearer",
        "expires_in": 3600,
    }

    play_response = MagicMock()
    play_response.status_code = 204

    monkeypatch.setenv("SPOTIFY_CLIENT_ID", "spotify-client-id")
    monkeypatch.setenv("SPOTIFY_CLIENT_SECRET", "spotify-client-secret")
    monkeypatch.setenv("SPOTIFY_REFRESH_TOKEN", "spotify-refresh-token")
    monkeypatch.setenv("SPOTIFY_DEVICE_ID", "device-from-env")
    monkeypatch.delenv("SPOTIFY_ACCESS_TOKEN", raising=False)

    with patch("src.skills.builtin.spotify_player_skill.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__.return_value = mock_client
        mock_client.__aexit__.return_value = None
        mock_client.post = AsyncMock(return_value=token_response)
        mock_client.put = AsyncMock(return_value=play_response)
        mock_client_cls.return_value = mock_client

        result = await manager.execute_tool(
            "spotify_player.play_playlist",
            {
                "playlist_id": "https://open.spotify.com/playlist/0123456789ABCDEFGHIJKL?si=test",
                "offset_position": 2,
                "position_ms": 1500,
            },
        )

    assert result.success is True
    assert result.data["started"] is True
    assert result.data["playlist_id"] == "0123456789ABCDEFGHIJKL"
    assert result.data["device_id"] == "device-from-env"

    token_call = mock_client.post.call_args
    assert token_call is not None
    assert token_call.kwargs["data"] == {
        "grant_type": "refresh_token",
        "refresh_token": "spotify-refresh-token",
    }

    play_call = mock_client.put.call_args
    assert play_call is not None
    assert play_call.args[0].endswith("/me/player/play")
    assert play_call.kwargs["params"] == {"device_id": "device-from-env"}
    assert play_call.kwargs["headers"]["Authorization"] == "Bearer spotify-user-token"
    assert play_call.kwargs["json"] == {
        "context_uri": "spotify:playlist:0123456789ABCDEFGHIJKL",
        "offset": {"position": 2},
        "position_ms": 1500,
    }


@pytest.mark.asyncio
async def test_googleworkspace_cli_check_cli_returns_version_output(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 7,
            "stdout": "gws 0.8.0\n",
            "stderr": "",
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool("googleworkspace_cli.check_cli", {})

    assert result.success is True
    assert result.data["available"] is True
    assert result.data["version"] == "gws 0.8.0"


@pytest.mark.asyncio
async def test_googleworkspace_cli_run_gws_command_wraps_cmd_on_windows(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}
    executable_path = r"C:\Users\pc\AppData\Roaming\npm\gws.cmd"

    class DummyProcess:
        returncode = 0

        async def communicate(self) -> tuple[bytes, bytes]:
            return b'{"ok":true}', b""

    async def fake_create_subprocess_exec(*args: Any, **kwargs: Any) -> DummyProcess:
        captured["args"] = list(args)
        captured["kwargs"] = kwargs
        return DummyProcess()

    monkeypatch.setattr(gws_skill, "_resolve_gws_executable", lambda requested, working_dir: executable_path)
    monkeypatch.setattr(gws_skill.asyncio, "create_subprocess_exec", fake_create_subprocess_exec)

    result = await gws_skill._run_gws_command(
        ["gws", "--version"],
        timeout_seconds=5,
        cwd=Path("."),
        env_overrides={},
    )

    assert captured["args"][0].lower().endswith("cmd.exe")
    assert captured["args"][1:4] == ["/d", "/c", executable_path]
    assert captured["args"][4:] == ["--version"]
    assert result["parsed_json"] == {"ok": True}


@pytest.mark.asyncio
async def test_googleworkspace_cli_run_readonly_command_serializes_params(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    temp_base = _create_workspace_temp_dir()
    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        captured["timeout_seconds"] = timeout_seconds
        captured["cwd"] = str(cwd) if cwd is not None else None
        captured["env_overrides"] = dict(env_overrides or {})
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 12,
            "stdout": '{"files":[{"id":"abc123"}]}',
            "stderr": "",
            "parsed_json": {"files": [{"id": "abc123"}]},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    try:
        result = await manager.execute_tool(
            "googleworkspace_cli.run_readonly_command",
            {
                "service": "drive",
                "resource": "files",
                "action": "list",
                "params": {"pageSize": 5, "q": "name contains 'Q1'"},
                "dry_run": True,
            },
            context=SkillContext(scratch_dir=temp_base),
        )
        assert result.success is True
        assert captured["args"][:4] == ["gws", "drive", "files", "list"]
        assert "--params" in captured["args"]
        params_index = captured["args"].index("--params")
        params_payload = json.loads(captured["args"][params_index + 1])
        assert params_payload == {"pageSize": 5, "q": "name contains 'Q1'"}
        assert "--dry-run" in captured["args"]
        assert captured["env_overrides"].get("GOOGLE_WORKSPACE_CLI_CONFIG_DIR")
        assert result.data["response"]["parsed_json"]["files"][0]["id"] == "abc123"
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_googleworkspace_cli_list_available_commands_parses_help(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 8,
            "stdout": (
                "Google Workspace CLI\n\n"
                "Commands:\n"
                "  files      Manage Drive files\n"
                "  permissions  Manage file permissions\n"
                "Options:\n"
                "  -h, --help   Show help\n"
            ),
            "stderr": "",
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.list_available_commands",
        {"service": "drive"},
    )
    assert result.success is True
    assert captured["args"] == ["gws", "drive", "--help"]
    assert result.data["scope"] == "drive"
    assert result.data["command_names"] == ["files", "permissions"]
    assert result.data["command_count"] == 2
    assert "response" not in result.data


@pytest.mark.asyncio
async def test_googleworkspace_cli_allows_create_actions(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 9,
            "stdout": '{"presentationId":"pres_123"}',
            "stderr": "",
            "parsed_json": {"presentationId": "pres_123"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "slides",
            "resource": "presentations",
            "action": "create",
            "json_payload": {"title": "Q2 Review"},
        },
    )
    assert result.success is True
    assert captured["args"][:4] == ["gws", "slides", "presentations", "create"]
    assert "--json" in captured["args"]
    assert result.data["response"]["parsed_json"]["presentationId"] == "pres_123"


@pytest.mark.asyncio
async def test_googleworkspace_cli_slides_batch_update_builds_requests_from_slides(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 10,
            "stdout": '{"replies":[]}',
            "stderr": "",
            "parsed_json": {"replies": []},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)

    temp_base = _create_workspace_temp_dir()
    try:
        images_dir = temp_base / "images"
        images_dir.mkdir(parents=True, exist_ok=True)
        (images_dir / "overview.png").write_bytes(b"overview")

        manager = SkillManager.from_manifest_directory("src/skills/manifests")
        result = await manager.execute_tool(
            "googleworkspace_cli.slides_batch_update_presentation",
            {
                "presentation_id": "pres_123",
                "slides": [
                    {
                        "title": "Overview",
                        "bullets": ["Goal"],
                        "images": [{"path": "images/overview.png", "alt": "Overview chart"}],
                    }
                ],
                "image_url_prefix": "https://cdn.example.com/decks",
            },
            context=SkillContext(scratch_dir=temp_base),
        )

        assert result.success is True
        assert captured["args"][:4] == ["gws", "slides", "presentations", "batchUpdate"]
        params = json.loads(captured["args"][captured["args"].index("--params") + 1])
        payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
        assert params == {"presentationId": "pres_123"}
        assert len(payload["requests"]) == 8
        assert payload["requests"][0]["createSlide"]["slideLayoutReference"]["predefinedLayout"] == "BLANK"
        title_shape = next(req["createShape"] for req in payload["requests"] if req.get("createShape", {}).get("objectId") == "title_1")
        assert title_shape["elementProperties"]["pageObjectId"] == "slide_1"
        body_insert = next(req["insertText"] for req in payload["requests"] if req.get("insertText", {}).get("objectId") == "body_1")
        assert body_insert["text"] == "Goal"
        bullet_request = next(req["createParagraphBullets"] for req in payload["requests"] if "createParagraphBullets" in req)
        assert bullet_request["objectId"] == "body_1"
        image_request = next(
            req["createImage"] for req in payload["requests"] if "createImage" in req
        )
        assert image_request["url"] == "https://cdn.example.com/decks/images/overview.png"
        assert result.data["request_source"] == "slides"
        assert result.data["image_request_count"] == 1
        assert "Updated Google Slides presentation pres_123" in result.message
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_googleworkspace_cli_slides_create_presentation_from_markdown_runs_create_then_batch_update(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:4] == ["gws", "slides", "presentations", "create"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456"}',
                "stderr": "",
                "parsed_json": {"presentationId": "pres_456"},
            }
        if command[:4] == ["gws", "slides", "presentations", "get"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456","slides":[{"objectId":"initial_slide_1"}]}',
                "stderr": "",
                "parsed_json": {
                    "presentationId": "pres_456",
                    "slides": [{"objectId": "initial_slide_1"}],
                },
            }
        if command[:4] == ["gws", "slides", "presentations", "batchUpdate"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 12,
                "stdout": '{"replies":[]}',
                "stderr": "",
                "parsed_json": {"replies": []},
            }
        raise AssertionError(f"Unexpected gws command: {command}")

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)

    temp_base = _create_workspace_temp_dir()
    try:
        images_dir = temp_base / "images"
        images_dir.mkdir(parents=True, exist_ok=True)
        (images_dir / "overview-metrics.png").write_bytes(b"overview")

        manager = SkillManager.from_manifest_directory("src/skills/manifests")
        markdown = "# Q3 Product Roadmap\n\n## Overview Metrics\n- Revenue trend\n"
        result = await manager.execute_tool(
            "googleworkspace_cli.slides_create_presentation_from_markdown",
            {
                "markdown": markdown,
                "attach_scratch_images": True,
                "image_dir": "images",
                "image_match_mode": "title",
                "image_url_prefix": "https://cdn.example.com/decks",
            },
            context=SkillContext(scratch_dir=temp_base),
        )

        assert result.success is True
        assert calls[0][:4] == ["gws", "slides", "presentations", "create"]
        assert calls[1][:4] == ["gws", "slides", "presentations", "get"]
        assert calls[2][:4] == ["gws", "slides", "presentations", "batchUpdate"]
        create_payload = json.loads(calls[0][calls[0].index("--json") + 1])
        batch_payload = json.loads(calls[2][calls[2].index("--json") + 1])
        assert create_payload == {"title": "Q3 Product Roadmap"}
        assert batch_payload["requests"][-1] == {"deleteObject": {"objectId": "initial_slide_1"}}
        assert batch_payload["requests"][0]["createSlide"]["slideLayoutReference"]["predefinedLayout"] == "BLANK"
        body_insert = next(
            req["insertText"] for req in batch_payload["requests"] if req.get("insertText", {}).get("objectId") == "body_2"
        )
        assert body_insert["text"] == "Revenue trend"
        bullet_request = next(
            req["createParagraphBullets"] for req in batch_payload["requests"] if "createParagraphBullets" in req
        )
        assert bullet_request["objectId"] == "body_2"
        image_request = next(
            req["createImage"] for req in batch_payload["requests"] if "createImage" in req
        )
        assert image_request["url"] == "https://cdn.example.com/decks/images/overview-metrics.png"
        assert result.data["presentation_id"] == "pres_456"
        assert result.data["title"] == "Q3 Product Roadmap"
        assert result.data["auto_attached_images"] == 1
        assert result.data["image_request_count"] == 1
        assert result.data["deleted_initial_slide"] is True
        assert result.data["deleted_initial_slide_id"] == "initial_slide_1"
        assert result.data["batch_update_retried_without_delete"] is False
        assert result.data["result_path"] == "presentations/q3-product-roadmap.slides-result.json"
        assert result.data["link_path"] == "presentations/q3-product-roadmap.slides-result.txt"
        manifest_path = temp_base / "presentations" / "q3-product-roadmap.slides-result.json"
        link_path = temp_base / "presentations" / "q3-product-roadmap.slides-result.txt"
        assert manifest_path.exists()
        assert link_path.exists()
        manifest_payload = json.loads(manifest_path.read_text(encoding="utf-8"))
        assert manifest_payload["presentation_id"] == "pres_456"
        assert manifest_payload["presentation_url"] == "https://docs.google.com/presentation/d/pres_456/edit"
        assert manifest_payload["deleted_initial_slide"] is True
        assert "Presentation URL: https://docs.google.com/presentation/d/pres_456/edit" in link_path.read_text(
            encoding="utf-8"
        )
        assert "Created Google Slides deck 'Q3 Product Roadmap'" in result.message
        assert "Removed the default blank opening slide." in result.message
        assert "Saved result manifest to presentations/q3-product-roadmap.slides-result.json" in result.message
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_googleworkspace_cli_slides_create_presentation_from_markdown_retries_without_delete_when_google_rejects_only_slide_deletion(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:4] == ["gws", "slides", "presentations", "create"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456"}',
                "stderr": "",
                "parsed_json": {"presentationId": "pres_456"},
            }
        if command[:4] == ["gws", "slides", "presentations", "get"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456","slides":[{"objectId":"initial_slide_1"}]}',
                "stderr": "",
                "parsed_json": {
                    "presentationId": "pres_456",
                    "slides": [{"objectId": "initial_slide_1"}],
                },
            }
        if command[:4] == ["gws", "slides", "presentations", "batchUpdate"]:
            payload = json.loads(command[command.index("--json") + 1])
            if payload["requests"] and payload["requests"][-1] == {"deleteObject": {"objectId": "initial_slide_1"}}:
                return {
                    "command": command,
                    "returncode": 0,
                    "duration_ms": 12,
                    "stdout": (
                        '{"error":{"code":400,"message":"Cannot delete object initial_slide_1: '
                        'presentations must have at least one slide"}}'
                    ),
                    "stderr": "",
                    "parsed_json": {
                        "error": {
                            "code": 400,
                            "message": (
                                "Cannot delete object initial_slide_1: presentations must have at least one slide"
                            ),
                        }
                    },
                }
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 12,
                "stdout": '{"replies":[]}',
                "stderr": "",
                "parsed_json": {"replies": []},
            }
        raise AssertionError(f"Unexpected gws command: {command}")

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)

    temp_base = _create_workspace_temp_dir()
    try:
        manager = SkillManager.from_manifest_directory("src/skills/manifests")
        result = await manager.execute_tool(
            "googleworkspace_cli.slides_create_presentation_from_markdown",
            {
                "markdown": "# Q3 Product Roadmap\n\n## Overview Metrics\n- Revenue trend\n",
            },
            context=SkillContext(scratch_dir=temp_base),
        )
        assert result.success is True
        batch_calls = [call for call in calls if call[:4] == ["gws", "slides", "presentations", "batchUpdate"]]
        assert len(batch_calls) == 2
        first_payload = json.loads(batch_calls[0][batch_calls[0].index("--json") + 1])
        second_payload = json.loads(batch_calls[1][batch_calls[1].index("--json") + 1])
        assert first_payload["requests"][-1] == {"deleteObject": {"objectId": "initial_slide_1"}}
        assert all(request != {"deleteObject": {"objectId": "initial_slide_1"}} for request in second_payload["requests"])
        assert result.data["deleted_initial_slide"] is False
        assert result.data["deleted_initial_slide_id"] == "initial_slide_1"
        assert result.data["batch_update_retried_without_delete"] is True
        assert "Kept the default opening slide" in result.message
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_googleworkspace_cli_slides_create_presentation_from_markdown_surfaces_batchupdate_json_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        if command[:4] == ["gws", "slides", "presentations", "create"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456"}',
                "stderr": "",
                "parsed_json": {"presentationId": "pres_456"},
            }
        if command[:4] == ["gws", "slides", "presentations", "get"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 11,
                "stdout": '{"presentationId":"pres_456","slides":[{"objectId":"initial_slide_1"}]}',
                "stderr": "",
                "parsed_json": {
                    "presentationId": "pres_456",
                    "slides": [{"objectId": "initial_slide_1"}],
                },
            }
        if command[:4] == ["gws", "slides", "presentations", "batchUpdate"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 12,
                "stdout": '{"error":{"code":400,"message":"Invalid requests[0].createSlide"}}',
                "stderr": "",
                "parsed_json": {"error": {"code": 400, "message": "Invalid requests[0].createSlide"}},
            }
        raise AssertionError(f"Unexpected gws command: {command}")

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)

    temp_base = _create_workspace_temp_dir()
    try:
        manager = SkillManager.from_manifest_directory("src/skills/manifests")
        result = await manager.execute_tool(
            "googleworkspace_cli.slides_create_presentation_from_markdown",
            {
                "markdown": "# Q3 Product Roadmap\n\n## Overview Metrics\n- Revenue trend\n",
            },
            context=SkillContext(scratch_dir=temp_base),
        )
        assert result.success is False
        assert "Invalid requests[0].createSlide" in result.message
    finally:
        shutil.rmtree(temp_base, ignore_errors=True)


@pytest.mark.asyncio
async def test_googleworkspace_cli_normalizes_gmail_compose_payload_for_draft_create(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 10,
            "stdout": '{"id":"draft_1"}',
            "stderr": "",
            "parsed_json": {"id": "draft_1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "drafts",
            "action": "create",
            "json_payload": {
                "to": "bob@example.com",
                "subject": "Test draft",
                "body": "Hello Bob",
            },
        },
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "drafts", "create"]
    assert "--json" in captured["args"]
    payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    assert isinstance(payload.get("message"), dict)
    raw_text = str(payload["message"].get("raw") or "")
    assert raw_text


@pytest.mark.asyncio
async def test_googleworkspace_cli_normalizes_gmail_compose_payload_for_message_send(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 9,
            "stdout": '{"id":"msg_sent"}',
            "stderr": "",
            "parsed_json": {"id": "msg_sent"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "messages",
            "action": "send",
            "json_payload": {
                "to": ["bob@example.com"],
                "subject": "Hello",
                "text": "Body text",
            },
        },
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "send"]
    assert "--json" in captured["args"]
    payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    assert "raw" in payload
    assert isinstance(payload["raw"], str) and len(payload["raw"]) > 20


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_list_unread_tool_builds_expected_query(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 9,
            "stdout": '{"messages":[]}',
            "stderr": "",
            "parsed_json": {"messages": []},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_list_unread",
        {"max_results": 5, "query": "from:billing@example.com"},
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["maxResults"] == 5
    assert "in:inbox is:unread" in params_payload["q"]
    assert "from:billing@example.com" in params_payload["q"]


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_list_all_tool_supports_pagination(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:5] == ["gws", "gmail", "users", "messages", "list"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 9,
                "stdout": (
                    '{"messages":[{"id":"m3","threadId":"t3"}],'
                    '"nextPageToken":"token-next","resultSizeEstimate":321}'
                ),
                "stderr": "",
                "parsed_json": {
                    "messages": [{"id": "m3", "threadId": "t3"}],
                    "nextPageToken": "token-next",
                    "resultSizeEstimate": 321,
                },
            }
        if command[:5] == ["gws", "gmail", "users", "messages", "get"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 7,
                "stdout": (
                    '{"id":"m3","threadId":"t3","snippet":"Preview all mail",'
                    '"payload":{"headers":[{"name":"From","value":"Carol <carol@example.com>"},'
                    '{"name":"Subject","value":"Quarterly update"},{"name":"Date","value":"Tue, 10 Mar 2026"}]},'
                    '"labelIds":["INBOX"]}'
                ),
                "stderr": "",
                "parsed_json": {
                    "id": "m3",
                    "threadId": "t3",
                    "snippet": "Preview all mail",
                    "payload": {
                        "headers": [
                            {"name": "From", "value": "Carol <carol@example.com>"},
                            {"name": "Subject", "value": "Quarterly update"},
                            {"name": "Date", "value": "Tue, 10 Mar 2026"},
                        ]
                    },
                    "labelIds": ["INBOX"],
                },
            }
        return {"command": command, "returncode": 0, "duration_ms": 6, "stdout": "{}", "stderr": "", "parsed_json": {}}

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_list_all",
        {
            "max_results": 7,
            "page_token": "token-prev",
            "query": "label:important",
        },
    )
    assert result.success is True
    params_payload = json.loads(calls[0][calls[0].index("--params") + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["maxResults"] == 7
    assert params_payload["pageToken"] == "token-prev"
    assert params_payload["q"] == "label:important"
    data = result.data or {}
    assert data["page_token"] == "token-prev"
    assert data["next_page_token"] == "token-next"
    assert data["result_size_estimate"] == 321
    assert data["message_count"] == 1
    summaries = data.get("gmail_message_summaries")
    assert isinstance(summaries, list) and len(summaries) == 1
    assert summaries[0]["subject"] == "Quarterly update"
    assert summaries[0]["from"] == "Carol <carol@example.com>"
    assert any(call[:5] == ["gws", "gmail", "users", "messages", "get"] for call in calls)


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_list_all_returns_up_to_requested_ten_summaries(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []
    list_messages = [{"id": f"m{i}", "threadId": f"t{i}"} for i in range(1, 11)]

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:5] == ["gws", "gmail", "users", "messages", "list"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 9,
                "stdout": json.dumps({"messages": list_messages}),
                "stderr": "",
                "parsed_json": {"messages": list_messages},
            }
        if command[:5] == ["gws", "gmail", "users", "messages", "get"]:
            params_payload = json.loads(command[command.index("--params") + 1])
            message_id = str(params_payload["id"])
            index = int(message_id.removeprefix("m"))
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 7,
                "stdout": json.dumps(
                    {
                        "id": message_id,
                        "threadId": f"t{index}",
                        "snippet": f"Preview {index}",
                        "payload": {
                            "headers": [
                                {"name": "From", "value": f"Sender {index} <sender{index}@example.com>"},
                                {"name": "Subject", "value": f"Subject {index}"},
                                {"name": "Date", "value": f"Tue, {index:02d} Mar 2026"},
                            ]
                        },
                    }
                ),
                "stderr": "",
                "parsed_json": {
                    "id": message_id,
                    "threadId": f"t{index}",
                    "snippet": f"Preview {index}",
                    "payload": {
                        "headers": [
                            {"name": "From", "value": f"Sender {index} <sender{index}@example.com>"},
                            {"name": "Subject", "value": f"Subject {index}"},
                            {"name": "Date", "value": f"Tue, {index:02d} Mar 2026"},
                        ]
                    },
                },
            }
        return {"command": command, "returncode": 0, "duration_ms": 6, "stdout": "{}", "stderr": "", "parsed_json": {}}

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_list_all",
        {"max_results": 10},
    )

    assert result.success is True
    data = result.data or {}
    assert data["max_results"] == 10
    assert data["message_count"] == 10
    summaries = data.get("gmail_message_summaries")
    assert isinstance(summaries, list) and len(summaries) == 10
    assert summaries[0]["subject"] == "Subject 1"
    assert summaries[-1]["subject"] == "Subject 10"
    get_calls = [call for call in calls if call[:5] == ["gws", "gmail", "users", "messages", "get"]]
    assert len(get_calls) == 10


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_get_message_tool_defaults_to_full(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 8,
            "stdout": '{"id":"m1"}',
            "stderr": "",
            "parsed_json": {"id": "m1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_get_message",
        {"message_id": "m1"},
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "get"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["id"] == "m1"
    assert params_payload["format"] == "full"


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_compose_draft_tool_builds_raw_message(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 11,
            "stdout": '{"id":"draft_1"}',
            "stderr": "",
            "parsed_json": {"id": "draft_1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_compose_draft",
        {
            "to": "bob@example.com",
            "subject": "Draft subject",
            "body_text": "Draft body",
        },
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "drafts", "create"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["userId"] == "me"
    json_payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    raw_text = str((json_payload.get("message") or {}).get("raw") or "")
    assert raw_text
    padded = raw_text + ("=" * (-len(raw_text) % 4))
    decoded = base64.urlsafe_b64decode(padded).decode("utf-8", errors="ignore")
    assert "Draft subject" in decoded
    assert "Draft body" in decoded


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_send_message_tool_builds_raw_message(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 10,
            "stdout": '{"id":"msg_1"}',
            "stderr": "",
            "parsed_json": {"id": "msg_1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_send_message",
        {
            "to": "alice@example.com",
            "subject": "Send subject",
            "body_text": "Send body",
        },
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "send"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["userId"] == "me"
    json_payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    raw_text = str(json_payload.get("raw") or "")
    assert raw_text
    padded = raw_text + ("=" * (-len(raw_text) % 4))
    decoded = base64.urlsafe_b64decode(padded).decode("utf-8", errors="ignore")
    assert "Send subject" in decoded
    assert "Send body" in decoded


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_mark_read_tool_calls_modify(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 8,
            "stdout": '{"id":"m1","labelIds":["INBOX"]}',
            "stderr": "",
            "parsed_json": {"id": "m1", "labelIds": ["INBOX"]},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")
    result = await manager.execute_tool(
        "googleworkspace_cli.gmail_mark_read",
        {"message_id": "m1"},
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "modify"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["id"] == "m1"
    json_payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    assert json_payload["removeLabelIds"] == ["UNREAD"]


@pytest.mark.asyncio
async def test_googleworkspace_cli_calendar_create_event_tool_builds_insert_command(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        captured["timeout_seconds"] = timeout_seconds
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 9,
            "stdout": '{"id":"evt_1"}',
            "stderr": "",
            "parsed_json": {"id": "evt_1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.calendar_create_event",
        {
            "summary": "Team sync",
            "start_time": "2026-03-12T09:00:00+11:00",
            "end_time": "2026-03-12T09:30:00+11:00",
            "calendar_id": "primary",
            "location": "Meeting Room 1",
            "description": "Weekly team sync",
            "attendees": ["alice@example.com", "bob@example.com"],
        },
    )
    assert result.success is True
    assert captured["args"][:3] == ["gws", "calendar", "+insert"]
    assert "--calendar" in captured["args"]
    assert captured["args"][captured["args"].index("--calendar") + 1] == "primary"
    assert captured["args"][captured["args"].index("--summary") + 1] == "Team sync"
    assert captured["args"][captured["args"].index("--start") + 1] == "2026-03-12T09:00:00+11:00"
    assert captured["args"][captured["args"].index("--end") + 1] == "2026-03-12T09:30:00+11:00"
    attendees = [
        captured["args"][index + 1]
        for index, token in enumerate(captured["args"])
        if token == "--attendee"
    ]
    assert attendees == ["alice@example.com", "bob@example.com"]
    assert result.data["response"]["parsed_json"]["id"] == "evt_1"


@pytest.mark.asyncio
async def test_googleworkspace_cli_calendar_cancel_event_tool_builds_delete_command(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 8,
            "stdout": "{}",
            "stderr": "",
            "parsed_json": {},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.calendar_cancel_event",
        {
            "event_id": "evt_123",
            "calendar_id": "work@example.com",
            "send_updates": "externalOnly",
        },
    )
    assert result.success is True
    assert captured["args"][:4] == ["gws", "calendar", "events", "delete"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload == {
        "calendarId": "work@example.com",
        "eventId": "evt_123",
        "sendUpdates": "externalOnly",
    }


@pytest.mark.asyncio
async def test_googleworkspace_cli_calendar_list_today_tool_uses_agenda_helper(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 7,
            "stdout": '{"events":[{"id":"evt_today","summary":"Standup","start":{"dateTime":"2026-03-10T09:00:00+11:00"},"end":{"dateTime":"2026-03-10T09:15:00+11:00"}}]}',
            "stderr": "",
            "parsed_json": {
                "events": [
                    {
                        "id": "evt_today",
                        "summary": "Standup",
                        "start": {"dateTime": "2026-03-10T09:00:00+11:00"},
                        "end": {"dateTime": "2026-03-10T09:15:00+11:00"},
                    }
                ]
            },
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.calendar_list_today",
        {"calendar": "Work"},
    )
    assert result.success is True
    assert captured["args"] == ["gws", "calendar", "+agenda", "--today", "--calendar", "Work"]
    summaries = (result.data or {}).get("calendar_event_summaries")
    assert isinstance(summaries, list) and len(summaries) == 1
    assert summaries[0]["summary"] == "Standup"
    assert summaries[0]["start"] == "2026-03-10T09:00:00+11:00"


@pytest.mark.asyncio
async def test_googleworkspace_cli_calendar_list_week_tool_extracts_nested_event_summaries(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 8,
            "stdout": '{"calendars":[{"calendarId":"primary","calendarSummary":"Work","events":[{"id":"evt_week","summary":"Planning","start":{"dateTime":"2026-03-11T10:00:00+11:00"},"end":{"dateTime":"2026-03-11T11:00:00+11:00"}}]}]}',
            "stderr": "",
            "parsed_json": {
                "calendars": [
                    {
                        "calendarId": "primary",
                        "calendarSummary": "Work",
                        "events": [
                            {
                                "id": "evt_week",
                                "summary": "Planning",
                                "start": {"dateTime": "2026-03-11T10:00:00+11:00"},
                                "end": {"dateTime": "2026-03-11T11:00:00+11:00"},
                            }
                        ],
                    }
                ]
            },
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.calendar_list_week",
        {},
    )
    assert result.success is True
    assert captured["args"] == ["gws", "calendar", "+agenda", "--week"]
    summaries = (result.data or {}).get("calendar_event_summaries")
    assert isinstance(summaries, list) and len(summaries) == 1
    assert summaries[0]["summary"] == "Planning"
    assert summaries[0]["calendar_id"] == "primary"
    assert summaries[0]["calendar_name"] == "Work"


@pytest.mark.asyncio
async def test_googleworkspace_cli_normalizes_legacy_gmail_draft_payload_shape(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 12,
            "stdout": '{"id":"draft_legacy_1"}',
            "stderr": "",
            "parsed_json": {"id": "draft_legacy_1"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "drafts",
            "action": "create",
            "json_payload": {
                "draft": {
                    "message": {
                        "payload": {
                            "headers": [
                                {"name": "To", "value": "bob@example.com"},
                                {"name": "Subject", "value": "Legacy draft"},
                            ],
                            "body": {"data": "I want to draft a test email"},
                        }
                    }
                }
            },
        },
    )
    assert result.success is True
    payload = json.loads(captured["args"][captured["args"].index("--json") + 1])
    assert isinstance(payload.get("message"), dict)
    raw_text = str(payload["message"].get("raw") or "")
    assert raw_text
    padded = raw_text + ("=" * (-len(raw_text) % 4))
    decoded = base64.urlsafe_b64decode(padded).decode("utf-8", errors="ignore")
    assert "I want to draft a test email" in decoded
    assert "Subject: Legacy draft" in decoded
    assert "To: bob@example.com" in decoded


@pytest.mark.asyncio
async def test_googleworkspace_cli_check_auth_requires_active_credentials(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 11,
            "stdout": json.dumps(
                {
                    "auth_method": "none",
                    "credential_source": "none",
                    "encrypted_credentials_exists": False,
                    "plain_credentials_exists": False,
                    "has_refresh_token": False,
                }
            ),
            "stderr": "",
            "parsed_json": {
                "auth_method": "none",
                "credential_source": "none",
                "encrypted_credentials_exists": False,
                "plain_credentials_exists": False,
                "has_refresh_token": False,
            },
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool("googleworkspace_cli.check_auth", {})
    assert result.success is False
    assert result.error_code == "execution_error"
    assert "no active credentials" in str(result.message).lower()


@pytest.mark.asyncio
async def test_googleworkspace_cli_run_readonly_command_supports_nested_and_helper_paths(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 10,
            "stdout": "{}",
            "stderr": "",
            "parsed_json": {},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    nested = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "users/messages",
            "action": "list",
            "params": {"userId": "me", "maxResults": 5},
        },
    )
    assert nested.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]

    shorthand = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "messages",
            "action": "list",
            "params": {"userId": "me", "maxResults": 1},
        },
    )
    assert shorthand.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]

    action_shorthand = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "action": "messages list",
            "params": {"userId": "me", "maxResults": 3},
        },
    )
    assert action_shorthand.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]

    default_resource = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "action": "list",
            "params": {"maxResults": 2},
        },
    )
    assert default_resource.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]
    assert "--params" in captured["args"]
    params_index = captured["args"].index("--params")
    params_payload = json.loads(captured["args"][params_index + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["maxResults"] == 2

    implicit_defaults = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "action": "list",
        },
    )
    assert implicit_defaults.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "list"]
    assert "--params" in captured["args"]
    params_index = captured["args"].index("--params")
    params_payload = json.loads(captured["args"][params_index + 1])
    assert params_payload["userId"] == "me"
    assert params_payload["maxResults"] == 10

    helper = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "action": "+triage",
        },
    )
    assert helper.success is True
    assert captured["args"][:3] == ["gws", "gmail", "+triage"]


@pytest.mark.asyncio
async def test_googleworkspace_cli_enriches_gmail_list_with_message_summaries(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:5] == ["gws", "gmail", "users", "messages", "list"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 8,
                "stdout": '{"messages":[{"id":"m1","threadId":"t1"}]}',
                "stderr": "",
                "parsed_json": {"messages": [{"id": "m1", "threadId": "t1"}]},
            }
        if command[:5] == ["gws", "gmail", "users", "messages", "get"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 7,
                "stdout": (
                    '{"id":"m1","threadId":"t1","snippet":"Preview text",'
                    '"payload":{"headers":[{"name":"From","value":"Alice <alice@example.com>"},'
                    '{"name":"Subject","value":"Status update"},{"name":"Date","value":"Mon, 9 Mar 2026"}]},'
                    '"labelIds":["INBOX","UNREAD"]}'
                ),
                "stderr": "",
                "parsed_json": {
                    "id": "m1",
                    "threadId": "t1",
                    "snippet": "Preview text",
                    "payload": {
                        "headers": [
                            {"name": "From", "value": "Alice <alice@example.com>"},
                            {"name": "Subject", "value": "Status update"},
                            {"name": "Date", "value": "Mon, 9 Mar 2026"},
                        ]
                    },
                    "labelIds": ["INBOX", "UNREAD"],
                },
            }
        return {"command": command, "returncode": 0, "duration_ms": 6, "stdout": "{}", "stderr": "", "parsed_json": {}}

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "messages",
            "action": "list",
            "params": {"maxResults": 1},
        },
    )
    assert result.success is True
    summaries = (result.data or {}).get("gmail_message_summaries")
    assert isinstance(summaries, list) and len(summaries) == 1
    assert summaries[0]["subject"] == "Status update"
    assert summaries[0]["from"] == "Alice <alice@example.com>"
    assert summaries[0]["snippet"] == "Preview text"
    assert any(call[:5] == ["gws", "gmail", "users", "messages", "get"] for call in calls)


@pytest.mark.asyncio
async def test_googleworkspace_cli_gmail_summary_falls_back_to_full_format_for_headers(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    calls: list[list[str]] = []

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        command = list(args)
        calls.append(command)
        if command[:5] == ["gws", "gmail", "users", "messages", "list"]:
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 8,
                "stdout": '{"messages":[{"id":"m2","threadId":"t2"}]}',
                "stderr": "",
                "parsed_json": {"messages": [{"id": "m2", "threadId": "t2"}]},
            }
        if command[:5] == ["gws", "gmail", "users", "messages", "get"]:
            params_payload = json.loads(command[command.index("--params") + 1])
            if params_payload.get("format") == "metadata":
                # Simulate metadata response without useful headers.
                return {
                    "command": command,
                    "returncode": 0,
                    "duration_ms": 7,
                    "stdout": '{"id":"m2","threadId":"t2","payload":{"headers":[]},"snippet":"Fallback preview"}',
                    "stderr": "",
                    "parsed_json": {
                        "id": "m2",
                        "threadId": "t2",
                        "payload": {"headers": []},
                        "snippet": "Fallback preview",
                    },
                }
            # full format fallback returns headers
            return {
                "command": command,
                "returncode": 0,
                "duration_ms": 7,
                "stdout": (
                    '{"id":"m2","threadId":"t2","snippet":"Fallback preview",'
                    '"payload":{"headers":[{"name":"Sender","value":"Bob <bob@example.com>"},'
                    '{"name":"Subject","value":"Invoice ready"},{"name":"Date","value":"Mon, 9 Mar 2026"}]}}'
                ),
                "stderr": "",
                "parsed_json": {
                    "id": "m2",
                    "threadId": "t2",
                    "snippet": "Fallback preview",
                    "payload": {
                        "headers": [
                            {"name": "Sender", "value": "Bob <bob@example.com>"},
                            {"name": "Subject", "value": "Invoice ready"},
                            {"name": "Date", "value": "Mon, 9 Mar 2026"},
                        ]
                    },
                },
            }
        return {"command": command, "returncode": 0, "duration_ms": 6, "stdout": "{}", "stderr": "", "parsed_json": {}}

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "messages",
            "action": "list",
            "params": {"maxResults": 1},
        },
    )
    assert result.success is True
    summaries = (result.data or {}).get("gmail_message_summaries")
    assert isinstance(summaries, list) and len(summaries) == 1
    assert summaries[0]["subject"] == "Invoice ready"
    assert summaries[0]["from"] == "Bob <bob@example.com>"
    assert summaries[0]["snippet"] == "Fallback preview"
    get_calls = [call for call in calls if call[:5] == ["gws", "gmail", "users", "messages", "get"]]
    assert len(get_calls) >= 2


@pytest.mark.asyncio
async def test_googleworkspace_cli_list_available_commands_normalizes_gmail_resource(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 6,
            "stdout": "Commands:\n  list  List messages\n",
            "stderr": "",
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.list_available_commands",
        {"service": "gmail", "resource": "messages"},
    )
    assert result.success is True
    assert captured["args"] == ["gws", "gmail", "users", "messages", "--help"]


@pytest.mark.asyncio
async def test_googleworkspace_cli_defaults_gmail_message_get_to_full_format(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.builtin import googleworkspace_cli_skill as gws_skill

    captured: Dict[str, Any] = {}

    async def fake_run_gws_command(
        args: Sequence[str],
        *,
        timeout_seconds: float,
        cwd: Path | None = None,
        env_overrides: Dict[str, str] | None = None,
    ) -> Dict[str, Any]:
        captured["args"] = list(args)
        return {
            "command": list(args),
            "returncode": 0,
            "duration_ms": 6,
            "stdout": '{"id":"m_full"}',
            "stderr": "",
            "parsed_json": {"id": "m_full"},
        }

    monkeypatch.setattr(gws_skill, "_run_gws_command", fake_run_gws_command)
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    result = await manager.execute_tool(
        "googleworkspace_cli.run_readonly_command",
        {
            "service": "gmail",
            "resource": "messages",
            "action": "get",
            "params": {"id": "abc123"},
        },
    )
    assert result.success is True
    assert captured["args"][:5] == ["gws", "gmail", "users", "messages", "get"]
    params_payload = json.loads(captured["args"][captured["args"].index("--params") + 1])
    assert params_payload["id"] == "abc123"
    assert params_payload["userId"] == "me"
    assert params_payload["format"] == "full"


@pytest.mark.asyncio
async def test_googleworkspace_cli_blocks_mutating_actions() -> None:
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    with pytest.raises(SkillValidationError, match="Destructive actions are blocked"):
        await manager.execute_tool(
            "googleworkspace_cli.run_readonly_command",
            {
                "service": "drive",
                "resource": "files",
                "action": "delete",
            },
            raise_errors=True,
        )


@pytest.mark.asyncio
async def test_github_project_manager_status_and_bump_with_mock_service(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.github import skill as gpm_skill

    class FakeGitHubIntegrationService:
        def __init__(self, workspace: Path) -> None:
            self.workspace = Path(workspace).resolve()

        @classmethod
        def from_env(
            cls, workspace: str | Path | None = None
        ) -> "FakeGitHubIntegrationService":
            return cls(Path(workspace or "."))

        def status(self) -> dict[str, Any]:
            return {
                "version": "1.2.3",
                "branch": "main",
                "ahead_by": 0,
                "behind_by": 0,
                "staged": [],
                "changed": [],
                "untracked": [],
            }

        def bump_version(self, bump: str) -> dict[str, Any]:
            return {
                "previous": "1.2.3",
                "current": "1.2.4" if bump == "patch" else "1.3.0",
                "bump": bump,
                "workspace": str(self.workspace),
            }

    monkeypatch.setattr(
        gpm_skill,
        "_load_integration_service_class",
        lambda: FakeGitHubIntegrationService,
    )
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    status_result = await manager.execute_tool("GitHubProjectManager.status", {"workspace": "."})
    assert status_result.success is True
    assert status_result.data["status"]["branch"] == "main"
    assert status_result.data["status"]["version"] == "1.2.3"

    bump_result = await manager.execute_tool(
        "GitHubProjectManager.bump_version",
        {"workspace": ".", "bump": "patch"},
    )
    assert bump_result.success is True
    assert bump_result.data["version"]["previous"] == "1.2.3"
    assert bump_result.data["version"]["current"] == "1.2.4"


@pytest.mark.asyncio
async def test_github_project_manager_extended_tools_with_mock_service(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from src.skills.github import skill as gpm_skill

    class FakeGitHubIntegrationService:
        def __init__(self, workspace: Path) -> None:
            self.workspace = Path(workspace).resolve()

        @classmethod
        def from_env(
            cls, workspace: str | Path | None = None
        ) -> "FakeGitHubIntegrationService":
            return cls(Path(workspace or "."))

        def fetch(self, remote_name: str | None = None) -> dict[str, Any]:
            return {"remote": remote_name or "origin", "branch": "main", "ahead_by": 0, "behind_by": 0}

        def pull(
            self,
            *,
            remote_name: str | None = None,
            branch: str | None = None,
            rebase: bool = False,
        ) -> dict[str, Any]:
            return {
                "remote": remote_name or "origin",
                "branch": branch or "main",
                "ahead_by": 0,
                "behind_by": 0,
                "rebase": rebase,
            }

        def push(
            self,
            *,
            remote_name: str | None = None,
            branch: str | None = None,
            set_upstream: bool = False,
            tags: bool = False,
        ) -> dict[str, Any]:
            return {
                "remote": remote_name or "origin",
                "branch": branch or "main",
                "set_upstream": set_upstream,
                "tags": tags,
            }

        def sync(
            self,
            *,
            remote_name: str | None = None,
            branch: str | None = None,
            rebase: bool = False,
            set_upstream: bool = False,
            tags: bool = False,
        ) -> dict[str, Any]:
            return {
                "pull": {"remote": remote_name or "origin", "branch": branch or "main", "rebase": rebase},
                "push": {"remote": remote_name or "origin", "branch": branch or "main", "set_upstream": set_upstream, "tags": tags},
            }

        def create_branch(
            self,
            branch: str,
            *,
            from_ref: str | None = None,
            push: bool = False,
            set_upstream: bool = True,
            remote_name: str | None = None,
        ) -> dict[str, Any]:
            return {
                "branch": branch,
                "created_from": from_ref,
                "push": {
                    "remote": remote_name or "origin",
                    "set_upstream": set_upstream,
                }
                if push
                else None,
            }

        def checkout_branch(self, branch: str) -> dict[str, Any]:
            return {"branch": branch, "ahead_by": 0, "behind_by": 0}

        def list_pull_requests(
            self,
            *,
            state: str = "open",
            sort: str = "created",
            direction: str = "desc",
            per_page: int = 30,
            page: int = 1,
        ) -> list[dict[str, Any]]:
            return [
                {
                    "number": 10,
                    "title": "Example",
                    "state": state,
                    "sort": sort,
                    "direction": direction,
                    "per_page": per_page,
                    "page": page,
                }
            ]

        def repository_info(self, *, include_rate_limit: bool = False) -> dict[str, Any]:
            payload: dict[str, Any] = {"full_name": "owner/CATBot", "has_rate_limit": include_rate_limit}
            if include_rate_limit:
                payload["rate_limit"] = {"remaining": 100}
            return payload

    monkeypatch.setattr(
        gpm_skill,
        "_load_integration_service_class",
        lambda: FakeGitHubIntegrationService,
    )
    manager = SkillManager.from_manifest_directory("src/skills/manifests")

    fetch_result = await manager.execute_tool("GitHubProjectManager.fetch", {"workspace": "."})
    pull_result = await manager.execute_tool("GitHubProjectManager.pull", {"workspace": ".", "rebase": True})
    push_result = await manager.execute_tool(
        "GitHubProjectManager.push",
        {"workspace": ".", "branch": "feature/test", "set_upstream": True, "tags": True},
    )
    sync_result = await manager.execute_tool("GitHubProjectManager.sync", {"workspace": ".", "branch": "main"})
    branch_result = await manager.execute_tool(
        "GitHubProjectManager.create_branch",
        {"workspace": ".", "branch": "feature/new", "from_ref": "main", "push": True},
    )
    checkout_result = await manager.execute_tool(
        "GitHubProjectManager.checkout_branch",
        {"workspace": ".", "branch": "main"},
    )
    list_pr_result = await manager.execute_tool(
        "GitHubProjectManager.list_pull_requests",
        {"workspace": ".", "state": "open", "per_page": 10, "page": 2},
    )
    repo_result = await manager.execute_tool(
        "GitHubProjectManager.repository_info",
        {"workspace": ".", "include_rate_limit": True},
    )

    assert fetch_result.success is True
    assert pull_result.success is True
    assert pull_result.data["pull"]["rebase"] is True
    assert push_result.success is True
    assert push_result.data["push"]["set_upstream"] is True
    assert sync_result.success is True
    assert sync_result.data["sync"]["pull"]["branch"] == "main"
    assert branch_result.success is True
    assert branch_result.data["branch"]["branch"] == "feature/new"
    assert checkout_result.success is True
    assert checkout_result.data["checkout"]["branch"] == "main"
    assert list_pr_result.success is True
    assert list_pr_result.data["pull_requests"][0]["page"] == 2
    assert repo_result.success is True
    assert repo_result.data["repository"]["rate_limit"]["remaining"] == 100


def _create_temp_filesystem_manifest_dir(base: Path, root: Path) -> Path:
    manifest_dir = base / "manifests"
    manifest_dir.mkdir(parents=True, exist_ok=True)
    root.mkdir(parents=True, exist_ok=True)
    manifest_path = manifest_dir / "filesystem.skill.json"
    manifest_path.write_text(
        (
            "{\n"
            '  "name": "filesystem",\n'
            '  "module": "src.skills.builtin.filesystem_skill:create_skill",\n'
            '  "settings": {\n'
            f'    "root_dir": "{str(root).replace("\\", "/")}"\n'
            "  }\n"
            "}\n"
        ),
        encoding="utf-8",
    )
    return manifest_dir


def _create_workspace_temp_dir() -> Path:
    base = Path("scratch") / f"skills-test-{uuid.uuid4().hex}"
    base.mkdir(parents=True, exist_ok=True)
    return base
